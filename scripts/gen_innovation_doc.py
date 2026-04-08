"""Convert 银行AIAgent创新文档_完整版.md to .docx with real diagrams embedded"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re, os

DIAGRAMS = r'C:/ai/CmbCoworkAgent/docs/diagrams'
MD_PATH  = r'C:/ai/CmbCoworkAgent/docs/银行AIAgent框架层创新文档_申报版_v4_精简.md'
OUT_PATH = r'C:/Users/87624/Desktop/银行AIAgent框架层创新文档_申报版_v10.docx'

FIG_MAP = {
    '图1':  'fig1_architecture.png',
    '图2':  'fig2_sandbox_layers.png',
    '图3':  'fig3_ssrf_flow.png',
    '图4':  'fig4_memory_layers.png',
    '图5':  'fig5_bm25_retrieval.png',
    '图6':  'fig6_prompt_hook.png',
    '图7':  'fig7_skill_lifecycle.png',
    '图8':  'fig8_compare_bar.png',
    '图9':  'fig9_memory_accuracy.png',
    '图10': 'fig10_token_compare.png',
    '图11': 'fig11_dream_flow.png',
    '图12': 'fig12_memory_panel.png',
}

doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def ea(run, name='宋体'):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)


def add_heading(document, t, level):
    p = document.add_heading(t, level=level)
    for run in p.runs:
        run.font.name = '黑体'
        run.font.size = Pt({1:18,2:15,3:13,4:12}.get(level,12))
        ea(run, '黑体')


def add_normal(document, t):
    p = document.add_paragraph(t)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        ea(run)


def flush_code(document, lines):
    if not lines:
        return
    p = document.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    ea(run, 'Courier New')


def flush_table(document, rows):
    if len(rows) < 2:
        return
    header = [c.strip() for c in rows[0].strip('|').split('|')]
    cols = len(header)
    t = document.add_table(rows=1, cols=cols)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = '黑体'
                run.font.size = Pt(9)
                ea(run, '黑体')
    for row_str in rows[2:]:
        row_str = row_str.strip()
        if not row_str.startswith('|'):
            continue
        cells = [c.strip() for c in row_str.strip('|').split('|')]
        new_row = t.add_row().cells
        for i in range(min(cols, len(cells))):
            new_row[i].text = cells[i]
            for para in new_row[i].paragraphs:
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(9)
                    ea(run)
    document.add_paragraph()


def add_figure(document, fig_key, caption):
    img_file = FIG_MAP.get(fig_key)
    img_path = os.path.join(DIAGRAMS, img_file) if img_file else None
    if img_path and os.path.exists(img_path):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.9))
    cap_p = document.add_paragraph(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_p.runs:
        run.font.name = '黑体'
        run.font.size = Pt(10)
        run.bold = True
        ea(run, '黑体')
    document.add_paragraph()


with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_code  = False
code_buf = []
table_buf = []

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # code fence
    if line.startswith('```'):
        if not in_code:
            in_code = True
            code_buf = []
        else:
            is_ascii = any(c in ''.join(code_buf) for c in ['┌','├','│','└','↓','▼','→','─'])
            if not is_ascii:
                flush_code(doc, code_buf)
            code_buf = []
            in_code = False
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # table
    if line.startswith('|'):
        table_buf.append(line)
        i += 1
        continue
    else:
        if table_buf:
            flush_table(doc, table_buf)
            table_buf = []

    # figure caption → insert image
    stripped = line.strip().lstrip('*').rstrip('*').strip()
    fig_match = re.match(r'^(图\d+)\s+(.+)$', stripped)
    if fig_match:
        fig_key = fig_match.group(1)
        if fig_key in FIG_MAP:
            add_figure(doc, fig_key, stripped)
            i += 1
            continue

    # headings
    if line.startswith('# '):
        add_heading(doc, line[2:], 1)
    elif line.startswith('## '):
        add_heading(doc, line[3:], 2)
    elif line.startswith('### '):
        add_heading(doc, line[4:], 3)
    elif line.startswith('#### '):
        add_heading(doc, line[5:], 4)
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        p.paragraph_format.left_indent = Cm(1)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            run.font.italic = True
            run.font.size = Pt(9)
    elif re.match(r'^\d+\. ', line):
        p = doc.add_paragraph(line, style='List Number')
        for run in p.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            ea(run)
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        for run in p.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            ea(run)
    elif line == '---':
        doc.add_paragraph()
    elif line.strip() == '':
        pass
    else:
        add_normal(doc, line)

    i += 1

if table_buf:
    flush_table(doc, table_buf)

doc.save(OUT_PATH)
print('Saved to', OUT_PATH)
